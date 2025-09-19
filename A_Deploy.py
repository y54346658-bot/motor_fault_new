# 在完成所有的功能的页面基础上，引入DeepSeek，优先调用DeepSeek对用户问题进行回答。
# 注意：KIMI API免费，DeepSeek API收费，需要申请API Key。

# 已成功通过RAG的方式加载故障手册，并使用CNN-1D模型进行故障诊断。

# 使用更轻量级的模型 all-MiniLM-L6-v2

# 在文件开头添加必要的导入
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
from sklearn.preprocessing import normalize as sklearn_normalize
import streamlit as st
import torch
import torch.nn as nn
import pandas as pd
import matplotlib.pyplot as plt
import os
import requests
import time
import re
import io
from datetime import datetime
from collections import Counter
import PyPDF2
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.platypus import ListFlowable, ListItem, Paragraph
from reportlab.lib.styles import ListStyle
import zipfile
import tempfile
from pathlib import Path
import shutil
import huggingface_hub
from huggingface_hub import hf_hub_download
from matplotlib import font_manager as fm
import matplotlib as mpl
import platform
import tempfile
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import inch
import os
import tempfile
import base64
import requests

# 获取当前日期和时间（可以根据需要调整时区，这里使用UTC+8北京时间为例）
current_time = datetime.now().strftime("%Y年%m月%d日 %H:%M")

# 设置中文字体支持
plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

# 解决中文显示问题的函数
def setup_chinese_font():
    """
    设置中文字体支持，适用于Streamlit Cloud环境
    优先使用项目中的字体文件，如果不可用则尝试系统字体
    """
    # 方法1: 尝试使用项目中的字体文件
    font_paths = [
        "./fonts/simhei.ttf"      # 黑体
    ]
    
    for font_path in font_paths:
        if os.path.exists(font_path):
            try:
                # 注册字体
                fm.fontManager.addfont(font_path)
                font_name = fm.FontProperties(fname=font_path).get_name()
                plt.rcParams['font.family'] = [font_name, 'sans-serif']
                plt.rcParams['axes.unicode_minus'] = False
                st.success(f"成功加载字体: {font_name}")
                return True
            except Exception as e:
                st.warning(f"字体加载失败 {font_path}: {e}")
    
    # 方法2: 尝试使用系统字体 (适用于Streamlit Cloud的Linux环境)
    system_fonts = [
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"
    ]
    
    for font_path in system_fonts:
        if os.path.exists(font_path):
            try:
                # 注册字体
                fm.fontManager.addfont(font_path)
                font_name = fm.FontProperties(fname=font_path).get_name()
                plt.rcParams['font.family'] = [font_name, 'sans-serif']
                plt.rcParams['axes.unicode_minus'] = False
                st.success(f"成功加载系统字体: {font_name}")
                return True
            except Exception as e:
                st.warning(f"系统字体加载失败 {font_path}: {e}")
    
    # 方法3: 尝试使用matplotlib内置的字体别名
    try:
        plt.rcParams['font.sans-serif'] = ['DejaVu Sans', 'Arial Unicode MS', 'sans-serif']
        plt.rcParams['axes.unicode_minus'] = False
        st.info("使用默认字体别名")
        return True
    except Exception as e:
        st.error(f"字体设置失败: {e}")
        return False

# 设置中文字体
if setup_chinese_font():
    st.success("中文字体配置成功")
else:
    st.error("中文字体配置失败，中文可能显示为方框")

# 定义CNN-1D模型
class CNN1DModel(nn.Module):
    def __init__(self, conv_archs, num_classes, batch_size, input_channels=1):
        super(CNN1DModel, self).__init__()
        self.batch_size = batch_size
        self.conv_arch = conv_archs
        self.input_channels = input_channels
        self.features = self._make_layers()
        self.avgpool = nn.AdaptiveAvgPool1d(9)
        self.classifier = nn.Sequential(
            nn.Linear(128 * 3 * 3, 500),
            nn.ReLU(inplace=True),
            nn.Dropout(),
            nn.Linear(500, num_classes),
        )
    
    def _make_layers(self):
        layers = []
        for (num_convs, out_channels) in self.conv_arch:
            for _ in range(num_convs):
                layers.append(nn.Conv1d(self.input_channels, out_channels, kernel_size=3, padding=1))
                layers.append(nn.ReLU(inplace=True))
                self.input_channels = out_channels
            layers.append(nn.MaxPool1d(kernel_size=2, stride=2))
        return nn.Sequential(*layers)

    def forward(self, input_seq):
        batch_size = input_seq.size(0)
        input_seq = input_seq.view(batch_size, 1, 1024)
        features = self.features(input_seq)
        x = self.avgpool(features)
        flat_tensor = x.view(batch_size, -1)
        output = self.classifier(flat_tensor)
        return output

# 设置页面
st.set_page_config(page_title="电机故障诊断系统", page_icon="🔧", layout="wide")

# 应用标题
st.title("🔧 电机故障诊断系统")

# 设备设置
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

# 定义故障类型映射
fault_type_mapping = {
    0: "正常状态 (de_normal)",
    1: "7mil内圈故障 (de_7_inner)",
    2: "7mil滚动体故障 (de_7_ball)",
    3: "7mil外圈故障 (de_7_outer)",
    4: "14mil内圈故障 (de_14_inner)",
    5: "14mil滚动体故障 (de_14_ball)",
    6: "14mil外圈故障 (de_14_outer)",
    7: "21mil内圈故障 (de_21_inner)",
    8: "21mil滚动体故障 (de_21_ball)",
    9: "21mil外圈故障 (de_21_outer)"
}

# 修改 load_model 函数
@st.cache_resource
def load_model():
    try:
        if not os.path.exists('best_model_cnn1d.pt'):
            st.error("模型文件不存在")
            return None
        
        # 添加 weights_only=False 参数
        model = torch.load('best_model_cnn1d.pt', map_location=device, weights_only=False)
        model.eval()
        return model
    except Exception as ex:
        st.error(f"模型加载失败: {ex}")
        return None

# 下载并加载 all-MiniLM-L6-v2 嵌入模型
# 确保目录存在
def ensure_directory_exists(path):
    """确保指定路径的目录存在，如果不存在则创建"""
    os.makedirs(path, exist_ok=True)
    return path

# 下载并加载 all-MiniLM-L6-v2 嵌入模型
@st.cache_resource(show_spinner=False)
def load_embedding_model():
    try:
        # 模型本地保存路径
        base_model_path = "./local_models_new"
        model_path = os.path.join(base_model_path, "all-MiniLM-L6-v2")
        pooling_dir = os.path.join(model_path, "1_Pooling")
        
        # 确保所有必要的目录都存在
        ensure_directory_exists(base_model_path)
        ensure_directory_exists(model_path)
        ensure_directory_exists(pooling_dir)
        
        # 检查模型是否已存在
        if not os.path.exists(os.path.join(model_path, "pytorch_model.bin")):
            with st.spinner("正在下载 all-MiniLM-L6-v2 模型，这可能需要几分钟..."):
                # 基础URL
                base_url = "https://hf-mirror.com/sentence-transformers/all-MiniLM-L6-v2/resolve/main"
                
                # 需要下载的文件列表
                files_to_download = [
                    "pytorch_model.bin",
                    "config.json",
                    "tokenizer.json",
                    "vocab.txt",
                    "special_tokens_map.json",
                    "tokenizer_config.json",
                    "sentence_bert_config.json",
                    "modules.json",
                    "1_Pooling/config.json"
                ]
                
                # 下载所有文件
                for filename in files_to_download:
                    # 构建完整的URL
                    file_url = f"{base_url}/{filename}"
                    
                    # 构建本地保存路径
                    if "/" in filename:
                        # 处理嵌套目录结构
                        dir_name = os.path.dirname(filename)
                        file_dir = os.path.join(model_path, dir_name)
                        ensure_directory_exists(file_dir)
                        filepath = os.path.join(model_path, filename)
                    else:
                        filepath = os.path.join(model_path, filename)
                    
                    # 下载文件
                    response = requests.get(file_url, stream=True)
                    if response.status_code == 200:
                        with open(filepath, 'wb') as f:
                            for chunk in response.iter_content(chunk_size=8192):
                                f.write(chunk)
                    else:
                        st.error(f"下载 {filename} 失败，状态码: {response.status_code}")
                        return None
        
        # 检查模型文件是否完整
        required_files = [
            "pytorch_model.bin",
            "config.json",
            "tokenizer.json",
            "vocab.txt",
            "special_tokens_map.json",
            "tokenizer_config.json",
            "sentence_bert_config.json",
            "modules.json",
            "1_Pooling/config.json"
        ]
        
        missing_files = []
        for file in required_files:
            if not os.path.exists(os.path.join(model_path, file)):
                missing_files.append(file)
        
        if missing_files:
            st.error(f"模型文件不完整，缺少: {', '.join(missing_files)}")
            # 尝试删除不完整的模型并重新下载
            shutil.rmtree(model_path, ignore_errors=True)
            return load_embedding_model()  # 递归调用自身重新下载
        
        # 加载本地模型
        model = SentenceTransformer(model_path)
        st.success("嵌入模型加载成功! (使用 all-MiniLM-L6-v2)")
        return model
    except Exception as e:
        st.error(f"嵌入模型加载失败: {e}")
        import traceback
        st.error(traceback.format_exc())
        # 尝试删除可能损坏的模型文件
        if 'model_path' in locals():
            shutil.rmtree(model_path, ignore_errors=True)
        return None

# 数据分割函数
def split_data_with_overlap(data, time_steps, label, overlap_ratio=0.5):
    stride = int(time_steps * (1 - overlap_ratio))
    samples = (len(data) - time_steps) // stride + 1
    data_list = []
    for i in range(samples):
        start_idx = i * stride
        end_idx = start_idx + time_steps
        temp_data = data[start_idx:end_idx].tolist()
        temp_data.append(label)
        data_list.append(temp_data)
    columns = list(range(time_steps)) + ['label']
    result_df = pd.DataFrame(data_list, columns=columns)
    return result_df

# 归一化函数 (重命名以避免与sklearn冲突)
def normalize_data(data):
    return (data - np.min(data)) / (np.max(data) - np.min(data))

# ==============================
# 优化后的向量知识库相关函数
# ==============================
def create_vector_knowledge_base_optimized(text, embedding_model, chunk_size=500, overlap=50, batch_size=16):
    """
    优化版的向量知识库创建函数
    :param text: 故障手册文本
    :param embedding_model: 嵌入模型
    :param chunk_size: 文本块大小
    :param overlap: 块间重叠大小
    :param batch_size: 批量处理大小
    :return: FAISS索引和文本块列表
    """
    # 检查参数有效性
    if overlap >= chunk_size:
        st.error("错误: 重叠大小必须小于文本块大小")
        return None, []
    
    if len(text) == 0:
        st.warning("文本内容为空，无法创建知识库")
        return None, []
    
    # 检查文本长度，如果过长则提醒用户
    if len(text) > 1000000:  # 约100万字符
        st.warning(f"文本长度较大 ({len(text)} 字符)，处理可能需要较长时间。建议使用较小的故障手册或增加块大小。")
    
    # 分割文本为块
    chunks = []
    start = 0
    
    # 如果文本长度小于块大小，直接使用整个文本
    if len(text) <= chunk_size:
        chunks.append(text)
    else:
        # 正常分块处理
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunk = text[start:end]
            chunks.append(chunk)
            
            # 如果已经到达文本末尾，则退出循环
            if end == len(text):
                break
                
            # 计算下一个起始位置，考虑重叠
            start = end - overlap
    
    st.info(f"已将文本分割为 {len(chunks)} 个文本块，开始生成嵌入向量...")
    
    # 显示进度
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    # 检查嵌入模型是否可用
    if embedding_model is None:
        st.warning("嵌入模型不可用，使用简单文本搜索模式")
        progress_bar.progress(1.0)
        status_text.text("完成!")
        return None, chunks
    
    try:
        # 批量生成嵌入向量（显著提高性能）
        embeddings = []
        total_batches = (len(chunks) + batch_size - 1) // batch_size
        
        # 添加中断按钮
        stop_processing = False
        if total_batches > 1:  # 只有在需要处理多个批次时才显示停止按钮
            if st.button("停止处理", key="stop_processing_btn"):
                stop_processing = True
        
        processed_count = 0
        
        for i in range(0, len(chunks), batch_size):
            if stop_processing:
                st.warning("用户中断了处理过程")
                return None, chunks[:i]  # 返回已处理的部分
            
            batch = chunks[i:i+batch_size]
            
            # 更新进度
            processed_count += len(batch)
            progress = min(processed_count / len(chunks), 1.0)
            progress_bar.progress(progress)
            status_text.text(f"处理中: {processed_count}/{len(chunks)} 个文本块")
            
            # 批量编码 - 添加更详细的错误处理
            try:
                batch_embeddings = embedding_model.encode(
                    batch, 
                    convert_to_tensor=False,
                    show_progress_bar=False,
                    batch_size=min(batch_size, len(batch)),
                    normalize_embeddings=True  # 直接归一化，避免后续步骤
                )
                embeddings.extend(batch_embeddings)
            except Exception as e:
                st.error(f"处理批次 {i//batch_size + 1}/{total_batches} 时出错: {e}")
                # 尝试处理较小的批次
                if len(batch) > 1:
                    st.info("尝试减小批量大小处理...")
                    try:
                        # 逐个处理
                        for single_chunk in batch:
                            single_embedding = embedding_model.encode(
                                [single_chunk],
                                convert_to_tensor=False,
                                show_progress_bar=False,
                                normalize_embeddings=True
                            )
                            embeddings.extend(single_embedding)
                    except Exception as e2:
                        st.error(f"单个处理也失败: {e2}")
                        # 跳过有问题的批次继续处理
                        continue
                else:
                    # 跳过有问题的批次继续处理
                    continue
        
        # 检查是否生成了任何嵌入向量
        if not embeddings:
            st.error("未能生成任何嵌入向量，所有处理批次都失败了")
            return None, chunks
            
        # 转换为numpy数组
        embeddings = np.array(embeddings)
        
        # 确保嵌入向量数量与文本块数量一致
        if len(embeddings) != len(chunks):
            st.warning(f"嵌入向量数量 ({len(embeddings)}) 与文本块数量 ({len(chunks)}) 不匹配")
            # 截取或填充以匹配数量
            min_len = min(len(embeddings), len(chunks))
            embeddings = embeddings[:min_len]
            chunks = chunks[:min_len]
        
        # 创建FAISS索引
        dimension = embeddings.shape[1]
        index = faiss.IndexFlatIP(dimension)  # 使用内积相似度
        
        # 添加向量到索引
        try:
            index.add(embeddings.astype('float32'))
        except Exception as e:
            st.error(f"创建FAISS索引失败: {e}")
            return None, chunks
        
        progress_bar.progress(1.0)
        status_text.text("完成!")
        
        # 显示内存使用情况
        if st.checkbox("显示内存使用情况", key="show_memory_usage"):
            index_size = index.ntotal * index.d * 4  # 大致计算索引大小(字节)
            st.write(f"FAISS索引大小: {index_size / (1024*1024):.2f} MB")
            st.write(f"文本块数量: {len(chunks)}")
            st.write(f"向量维度: {dimension}")
        
        return index, chunks
    except Exception as e:
        st.error(f"创建向量索引失败: {e}")
        import traceback
        st.error(f"详细错误信息: {traceback.format_exc()}")
        progress_bar.progress(1.0)
        status_text.text("出错!")
        return None, chunks

def preprocess_text(text):
    """
    预处理文本，移除多余空格和特殊字符
    :param text: 原始文本
    :return: 处理后的文本
    """
    if not text:
        return ""
    
    # 移除多余空格
    text = re.sub(r'\s+', ' ', text)
    # 移除特殊字符（保留中文、英文、数字和基本标点）
    text = re.sub(r'[^\w\u4e00-\u9fff\s.,!?;:，。！？；：]', '', text)
    return text.strip()

def search_vector_knowledge_base(query, index, chunks, embedding_model, top_k=3):
    """
    在向量知识库中搜索相关文档
    :param query: 查询文本
    :param index: FAISS索引
    :param chunks: 文本块列表
    :param embedding_model: 嵌入模型
    :param top_k: 返回最相关的k个结果
    :return: 相关文本块列表
    """
    # 如果索引不可用，使用简单文本搜索
    if index is None:
        return search_fallback_knowledge_base(query, chunks, top_k)
    
    try:
        # 生成查询向量
        query_embedding = embedding_model.encode([query], convert_to_tensor=False)
        query_embedding = sklearn_normalize(query_embedding, norm='l2', axis=1)
        
        # 搜索相似向量
        distances, indices = index.search(query_embedding.astype('float32'), top_k)
        
        # 获取相关文本块
        results = []
        for i, idx in enumerate(indices[0]):
            if idx < len(chunks):
                results.append({
                    "text": chunks[idx],
                    "score": distances[0][i]
                })
        
        return results
    except Exception as e:
        st.error(f"向量搜索失败: {e}")
        return search_fallback_knowledge_base(query, chunks, top_k)

def search_fallback_knowledge_base(query, chunks, top_k=3):
    """简单关键词匹配搜索"""
    results = []
    query_words = query.lower().split()
    
    for i, chunk in enumerate(chunks):
        score = 0
        chunk_lower = chunk.lower()
        for word in query_words:
            if word in chunk_lower:
                score += 1
        
        if score > 0:
            results.append({
                "text": chunk,
                "score": score / len(query_words)
            })
    
    # 按分数排序并返回前k个
    results.sort(key=lambda x: x["score"], reverse=True)
    return results[:top_k]

# ==============================
# API 相关函数
# ==============================
KIMI_API_URL = "https://api.moonshot.cn/v1/chat/completions"
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"

# 读取配置
def get_api_config(api_name):
    config = {}
    try:
        if api_name == "kimi":
            config["API_KEY"] = st.secrets["KIMI_API_KEY"]
        elif api_name == "deepseek":
            config["API_KEY"] = st.secrets["DEEPSEEK_API_KEY"]
    except (KeyError, FileNotFoundError):
        if api_name == "kimi":
            config["API_KEY"] = os.environ.get("KIMI_API_KEY", "")
        elif api_name == "deepseek":
            config["API_KEY"] = os.environ.get("DEEPSEEK_API_KEY", "")
    
    if not config["API_KEY"]:
        if api_name == "kimi":
            st.error("未找到KIMI API密钥，请在环境变量或secrets.toml中设置KIMI_API_KEY")
        elif api_name == "deepseek":
            st.error("未找到DeepSeek API密钥，请在环境变量或secrets.toml中设置DEEPSEEK_API_KEY")
    
    return config

# 从文件提取文本内容
def extract_text_from_file(file):
    """从上传的文件中提取文本内容"""
    file_type = file.type
    
    try:
        if file_type == "text/plain":
            # TXT文件
            return str(file.read(), "utf-8")
        
        elif file_type == "application/pdf":
            # PDF文件
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
        
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # DOCX文件
            doc = Document(file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        else:
            st.error("不支持的文件格式")
            return None
    except Exception as e:
        st.error(f"文件读取失败: {e}")
        return None

def call_deepseek_api(messages, max_tokens=2000, temperature=0.7):
    """调用DeepSeek API"""
    config = get_api_config("deepseek")
    
    # 确保 config 不为 None 且包含 API_KEY
    if not config or not config.get("API_KEY"):
        st.error("未找到DeepSeek API配置或API密钥，请检查配置")
        return None
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {config['API_KEY']}"
    }
    
    # 准备系统提示词
    system_prompt_content = prepare_system_prompt()
    
    system_prompt = {
        "role": "system", 
        "content": system_prompt_content
    }
    
    # 构建完整的消息列表，确保系统提示词在最前面
    api_messages = [system_prompt]
    
    # 添加所有非系统消息
    for msg in messages:
        if isinstance(msg, dict) and msg.get("role") != "system":
            api_messages.append({
                "role": msg.get("role", "user"),
                "content": msg.get("content", "")
            })
    
    data = {
        "model": "deepseek-chat",
        "messages": api_messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
        "stream": False
    }
    
    try:
        with st.spinner("正在调用DeepSeek API，请稍候..."):
            response = requests.post(
                DEEPSEEK_API_URL,
                headers=headers,
                json=data,
                timeout=60
            )
        
        # 检查响应状态
        if response.status_code == 200:
            try:
                result = response.json()
                # 检查响应结构
                if (result and isinstance(result, dict) and 
                    "choices" in result and isinstance(result["choices"], list) and 
                    len(result["choices"]) > 0 and 
                    "message" in result["choices"][0] and 
                    "content" in result["choices"][0]["message"]):
                    return result["choices"][0]["message"]["content"]
                else:
                    st.error(f"DeepSeek API响应格式异常: {result}")
                    return None
            except ValueError:
                st.error("DeepSeek API返回了无效的JSON响应")
                return None
        elif response.status_code == 401:
            st.error("DeepSeek API密钥无效或已过期，请检查API密钥")
            return None
        elif response.status_code == 429:
            st.error("DeepSeek API调用频率限制，请稍后再试")
            return None
        else:
            st.error(f"DeepSeek API请求失败，状态码: {response.status_code}, 响应: {response.text}")
            return None
            
    except requests.exceptions.Timeout:
        st.error("DeepSeek API请求超时，请稍后再试")
        return None
    except requests.exceptions.ConnectionError:
        st.error("网络连接错误，请检查网络设置")
        return None
    except Exception as err:
        st.error(f"DeepSeek API请求失败: {str(err)}")
        return None

def call_kimi_api(messages, max_tokens=2000, temperature=0.7):
    """调用KIMI API"""
    config = get_api_config("kimi")
    
    # 确保 config 不为 None 且包含 API_KEY
    if not config or not config.get("API_KEY"):
        st.error("未找到KIMI API配置或API密钥，请检查配置")
        return None
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {config['API_KEY']}"
    }
    
    # 准备系统提示词
    system_prompt_content = prepare_system_prompt()
    
    system_prompt = {
        "role": "system", 
        "content": system_prompt_content
    }
    
    # 构建完整的消息列表，确保系统提示词在最前面
    api_messages = [system_prompt]
    
    # 添加所有非系统消息
    for msg in messages:
        if isinstance(msg, dict) and msg.get("role") != "system":
            api_messages.append({
                "role": msg.get("role", "user"),
                "content": msg.get("content", "")
            })
    
    data = {
        "model": "moonshot-v1-8k",
        "messages": api_messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
        "top_p": 0.9
    }
    
    try:
        with st.spinner("正在调用KIMI API，请稍候..."):
            response = requests.post(
                KIMI_API_URL,
                headers=headers,
                json=data,
                timeout=60
            )
        
        # 检查响应状态
        if response.status_code == 200:
            try:
                result = response.json()
                # 检查响应结构
                if (result and isinstance(result, dict) and 
                    "choices" in result and isinstance(result["choices"], list) and 
                    len(result["choices"]) > 0 and 
                    "message" in result["choices"][0] and 
                    "content" in result["choices"][0]["message"]):
                    return result["choices"][0]["message"]["content"]
                else:
                    st.error(f"KIMI API响应格式异常: {result}")
                    return None
            except ValueError:
                st.error("KIMI API返回了无效的JSON响应")
                return None
        elif response.status_code == 401:
            st.error("KIMI API密钥无效或已过期，请检查API密钥")
            return None
        elif response.status_code == 429:
            st.error("KIMI API调用频率限制，请稍后再试")
            return None
        else:
            st.error(f"KIMI API请求失败，状态码: {response.status_code}, 响应: {response.text}")
            return None
            
    except requests.exceptions.Timeout:
        st.error("KIMI API请求超时，请稍后再试")
        return None
    except requests.exceptions.ConnectionError:
        st.error("网络连接错误，请检查网络设置")
        return None
    except Exception as err:
        st.error(f"KIMI API请求失败: {str(err)}")
        return None

def prepare_system_prompt():
    """准备系统提示词"""
    # 检查是否有诊断结果
    diagnosis_info = ""
    if 'diagnosis_results' in st.session_state and st.session_state['diagnosis_results'] is not None:
        diagnosis = st.session_state['diagnosis_results']
        # 确保 diagnosis 是字典类型
        if isinstance(diagnosis, dict):
            diagnosis_info = f"""
            已知振动信号经过CNN模型诊断，结果如下：
            - 诊断结果: {diagnosis.get('diagnosis_class', '未知')}
            - 置信度: {diagnosis.get('confidence_level', 0)*100:.2f}%
            - 信号统计: {diagnosis.get('analysis_results', '无')}
            """

            # 如果有报告内容，也添加到信息中
            if diagnosis.get('report'):
                diagnosis_info += f"\n- 详细报告: {diagnosis.get('report')[:200]}..."  # 只取前200字符避免过长
    
    # 准备故障手册信息 - 使用向量知识库检索相关内容
    manual_info = ""
    
    # 添加更严格的检查条件
    vector_index_exists = (
        'vector_index' in st.session_state and 
        st.session_state['vector_index'] is not None and
        hasattr(st.session_state['vector_index'], 'ntotal') and
        st.session_state['vector_index'].ntotal > 0
    )
    
    vector_chunks_exists = (
        'vector_chunks' in st.session_state and 
        st.session_state['vector_chunks'] is not None and
        len(st.session_state['vector_chunks']) > 0
    )
    
    if vector_index_exists and vector_chunks_exists:
        # 获取用户的最新查询
        user_query = ""
        for msg in reversed(st.session_state["messages"]):
            if msg.get("role") == "user":
                user_query = msg.get("content", "")
                break
        
        if user_query:
            # 从向量知识库中检索相关内容
            embedding_model = load_embedding_model()  # 获取嵌入模型实例来实时转换用户查询
            if embedding_model:
                relevant_chunks = search_vector_knowledge_base(
                    user_query, 
                    st.session_state['vector_index'],
                    st.session_state['vector_chunks'],
                    embedding_model,
                    top_k=3
                )
                
                if relevant_chunks:
                    manual_info = "以下是根据故障手册检索到的相关信息：\n"
                    for i, chunk in enumerate(relevant_chunks):
                        manual_info += f"\n[相关段落 {i+1}, 相似度: {chunk['score']:.3f}]\n{chunk['text']}\n"
    
    # 准备系统提示词
    system_prompt_content = f"""你是一名资深的电机故障诊断专家，熟悉振动信号分析、轴承/转子故障机理和维修方法。

    当前系统时间：{current_time}

    {manual_info}

    {diagnosis_info}

    你的任务是：根据用户的问题、故障手册参考和可用的诊断结果，提供专业、准确的解答和建议。

    请保持专业、严谨、工程化的语气，不要虚构不存在的数据。

    回答时请优先参考故障手册中的信息，并结合诊断结果给出建议。
    """
    
    return system_prompt_content

def call_ai_api(messages, max_tokens=2000, temperature=0.7):
    """根据用户选择的API偏好调用相应的API"""
    # 获取用户选择的API偏好
    preferred_api = st.session_state.get('preferred_api', 'deepseek')  # 默认为deepseek
    
    if preferred_api == "deepseek":
        # 先尝试调用DeepSeek API
        reply = call_deepseek_api(messages, max_tokens, temperature)
        
        # 如果DeepSeek调用失败，尝试调用KIMI API
        if reply is None:
            st.warning("DeepSeek API调用失败，尝试使用KIMI API")
            reply = call_kimi_api(messages, max_tokens, temperature)
    else:
        # 先尝试调用KIMI API
        reply = call_kimi_api(messages, max_tokens, temperature)
        
        # 如果KIMI调用失败，尝试调用DeepSeek API
        if reply is None:
            st.warning("KIMI API调用失败，尝试使用DeepSeek API")
            reply = call_deepseek_api(messages, max_tokens, temperature)
    
    return reply


# 处理 Markdown 格式的文本
def process_markdown_for_docx(text):
    """
    将Markdown格式文本处理为适合DOCX的格式
    :param text: Markdown格式的文本
    :return: 处理后的文本
    """
    # 首先规范化换行符
    text = re.sub(r'\r\n', '\n', text)  # 统一换行符
    text = re.sub(r'\n{3,}', '\n\n', text)  # 多个连续换行缩减为两个
    
    # 处理标题 - 确保标题后有额外的换行
    text = re.sub(r'^# (.*)$', r'标题1: \1\n\n', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.*)$', r'标题2: \1\n\n', text, flags=re.MULTILINE)
    text = re.sub(r'^### (.*)$', r'标题3: \1\n\n', text, flags=re.MULTILINE)
    text = re.sub(r'^#### (.*)$', r'标题4: \1\n\n', text, flags=re.MULTILINE)
    text = re.sub(r'^##### (.*)$', r'标题5: \1\n\n', text, flags=re.MULTILINE)
    text = re.sub(r'^###### (.*)$', r'标题6: \1\n\n', text, flags=re.MULTILINE)
    
    # 处理加粗 - 保留标记供后续处理
    text = re.sub(r'\*\*(.*?)\*\*', r'[BOLD]\1[/BOLD]', text)
    
    # 处理斜体 - 保留标记供后续处理
    text = re.sub(r'\*(.*?)\*', r'[ITALIC]\1[/ITALIC]', text)
    
    # 处理无序列表 - 确保列表项后有换行
    text = re.sub(r'^-\s+(.*)$', r'• \1\n', text, flags=re.MULTILINE)
    text = re.sub(r'^\*\s+(.*)$', r'• \1\n', text, flags=re.MULTILINE)
    
    # 处理有序列表 - 确保列表项后有换行
    def replace_ordered_list(match):
        return f"{match.group(1)}. {match.group(2)}\n"
    
    text = re.sub(r'^(\d+)\.\s+(.*)$', replace_ordered_list, text, flags=re.MULTILINE)
    
    # 处理分割线 - 使用特殊标记并添加额外换行
    text = re.sub(r'^---+\s*$', r'[HR]\n\n', text, flags=re.MULTILINE)
    text = re.sub(r'^\*\*\*+\s*$', r'[HR]\n\n', text, flags=re.MULTILINE)
    
    # 确保段落之间有足够的间距
    text = re.sub(r'\n\n+', '\n\n', text)
    
    # 处理换行 - 使用特殊标记
    text = text.replace('\n', '[NEWLINE]')
    
    return text

# 处理 Markdown 格式的文本
def process_markdown(text):
    """
    将 Markdown 格式转换为 ReportLab 可识别的格式
    """
    # 处理标题 (支持 1-6 级标题)
    text = re.sub(r'^# (.*)$', r'<b><font size="16">\1</font></b>', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.*)$', r'<b><font size="14">\1</font></b>', text, flags=re.MULTILINE)
    text = re.sub(r'^### (.*)$', r'<b><font size="12">\1</font></b>', text, flags=re.MULTILINE)
    text = re.sub(r'^#### (.*)$', r'<b><font size="11">\1</font></b>', text, flags=re.MULTILINE)
    text = re.sub(r'^##### (.*)$', r'<b><font size="10">\1</font></b>', text, flags=re.MULTILINE)
    text = re.sub(r'^###### (.*)$', r'<b><font size="9">\1</font></b>', text, flags=re.MULTILINE)
    
    # 处理加粗
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    
    # 处理斜体
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    
    # 处理无序列表 - 使用HTML列表标签
    text = re.sub(r'^-\s+(.*)$', r'<li>\1</li>', text, flags=re.MULTILINE)
    text = re.sub(r'^\*\s+(.*)$', r'<li>\1</li>', text, flags=re.MULTILINE)
    
    # 处理有序列表 - 使用HTML列表标签
    def replace_ordered_list(match):
        return f'<li value="{match.group(1)}">.{match.group(2)}</li>'
    
    text = re.sub(r'^(\d+)\.\s+(.*)$', replace_ordered_list, text, flags=re.MULTILINE)
    
    # 包裹列表项在ul或ol标签中
    lines = text.split('\n')
    in_list = False
    list_type = None  # 'ul' 或 'ol'
    processed_lines = []
    
    for line in lines:
        if line.startswith('<li>') or line.startswith('<li value='):
            if not in_list:
                # 确定列表类型
                if line.startswith('<li value='):
                    list_type = 'ol'
                    processed_lines.append(f'<{list_type}>')
                else:
                    list_type = 'ul'
                    processed_lines.append(f'<{list_type}>')
                in_list = True
            processed_lines.append(line)
        else:
            if in_list:
                processed_lines.append(f'</{list_type}>')
                in_list = False
                list_type = None
            processed_lines.append(line)
    
    # 处理最后可能还在列表中的情况
    if in_list:
        processed_lines.append(f'</{list_type}>')
    
    text = '\n'.join(processed_lines)
    
    # 处理分割线
    text = re.sub(r'^---+\s*$', r'<hr/>', text, flags=re.MULTILINE)
    text = re.sub(r'^\*\*\*+\s*$', r'<hr/>', text, flags=re.MULTILINE)
    
    # 处理换行
    text = text.replace('\n', '<br/>')
    
    return text

def create_docx_report(diagnosis_class, confidence_level, analysis_results, llm_report):
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体（尝试使用中文字体）
    try:
        # 设置全局字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'SimSun'  # 宋体
        font.size = Pt(12)
        # 设置西文字体
        font.name = 'Times New Roman'
        # 设置中文字体
        r = style._element.rPr
        rFonts = r.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), 'SimSun')
    except:
        # 如果设置失败，使用默认字体
        pass
    
    # 添加标题
    title = doc.add_heading('电机故障诊断报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    # 添加空行
    doc.add_paragraph()
    
    # 诊断信息
    subtitle = doc.add_heading('诊断结果', level=1)
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.bold = True
    
    diagnosis_para = doc.add_paragraph()
    diagnosis_para.add_run('故障类型: ').bold = True
    diagnosis_para.add_run(f'{diagnosis_class}')
    
    confidence_para = doc.add_paragraph()
    confidence_para.add_run('置信度: ').bold = True
    confidence_para.add_run(f'{confidence_level*100:.2f}%')
    
    # 添加空行
    doc.add_paragraph()
    
    # 分析结果
    analysis_title = doc.add_heading('信号分析', level=1)
    analysis_title_run = analysis_title.runs[0]
    analysis_title_run.font.size = Pt(14)
    analysis_title_run.font.bold = True
    
    # 创建分析结果表格
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # 设置表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '参数'
    hdr_cells[1].text = '值'
    
    # 设置表头格式
    for cell in hdr_cells:
        paragraph = cell.paragraphs[0]
        paragraph.runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加数据行
    for key, value in analysis_results.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)
    
    # 添加空行
    doc.add_paragraph()
    
    # 添加AI分析标题
    analysis_title = doc.add_heading('AI分析与建议', level=1)
    analysis_title_run = analysis_title.runs[0]
    analysis_title_run.font.size = Pt(14)
    analysis_title_run.font.bold = True
    
    # 处理LLM报告内容
    processed_report = process_markdown_for_docx(llm_report)
    
    # 分割处理后的报告为段落
    sections = processed_report.split('[NEWLINE][NEWLINE]')
    
    for section in sections:
        section = section.strip()
        if not section:
            continue
            
        # 处理分割线
        if section == '[HR]':
            # 添加分割线（使用段落边框）
            para = doc.add_paragraph()
            para_format = para.paragraph_format
            
            # 修复：使用正确的边框样式设置
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_UNDERLINE
            
            # 使用简单的方法添加分割线
            run = para.add_run("―" * 50)  # 使用长破折号作为分割线
            run.font.color.rgb = RGBColor(200, 200, 200)  # 灰色
            run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        # 处理标题
        if section.startswith('标题1: '):
            title = doc.add_heading(section[4:].strip(), level=1)
            title_run = title.runs[0]
            title_run.font.size = Pt(16)
            title_run.font.bold = True
        elif section.startswith('标题2: '):
            title = doc.add_heading(section[4:].strip(), level=2)
            title_run = title.runs[0]
            title_run.font.size = Pt(14)
            title_run.font.bold = True
        elif section.startswith('标题3: '):
            title = doc.add_heading(section[4:].strip(), level=3)
            title_run = title.runs[0]
            title_run.font.size = Pt(12)
            title_run.font.bold = True
        elif section.startswith('标题4: '):
            title = doc.add_heading(section[4:].strip(), level=4)
            title_run = title.runs[0]
            title_run.font.size = Pt(11)
            title_run.font.bold = True
        elif section.startswith('标题5: '):
            title = doc.add_heading(section[4:].strip(), level=5)
            title_run = title.runs[0]
            title_run.font.size = Pt(10)
            title_run.font.bold = True
        elif section.startswith('标题6: '):
            title = doc.add_heading(section[4:].strip(), level=6)
            title_run = title.runs[0]
            title_run.font.size = Pt(9)
            title_run.font.bold = True
        else:
            # 处理普通段落
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(10)
            
            # 处理加粗和斜体标记
            parts = re.split(r'(\[BOLD\].*?\[/BOLD\]|\[ITALIC\].*?\[/ITALIC\])', section)
            for part in parts:
                if part.startswith('[BOLD]') and part.endswith('[/BOLD]'):
                    run = paragraph.add_run(part[6:-7])
                    run.bold = True
                elif part.startswith('[ITALIC]') and part.endswith('[/ITALIC]'):
                    run = paragraph.add_run(part[8:-9])
                    run.italic = True
                else:
                    # 处理普通文本中的换行
                    sub_parts = part.split('[NEWLINE]')
                    for i, sub_part in enumerate(sub_parts):
                        if i > 0:
                            paragraph.add_run().add_break()  # 添加换行
                        paragraph.add_run(sub_part)
    
    # 添加页脚
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    footer_run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色
    footer_run.font.size = Pt(10)
    
    # 保存到字节流
    buffer = io.BytesIO()
    doc.save(buffer)
    
    # 获取字节数据
    docx_data = buffer.getvalue()
    buffer.close()
    
    return docx_data

# 创建PDF报告（使用ReportLab）- 针对Streamlit Cloud优化
def create_pdf_report(diagnosis_class, confidence_level, analysis_results, llm_report):
    # 创建字节缓冲区
    buffer = io.BytesIO()
    
    # 创建PDF文档
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    
    # 故事元素列表
    story = []
    
    # 获取样式
    styles = getSampleStyleSheet()
    
    # 尝试注册中文字体
    font_name = 'Helvetica'  # 默认字体
    
    try:
        # 方法1: 尝试使用项目中的字体文件
        font_paths = [
            "./fonts/simhei.ttf"
        ]
        
        font_registered = False
        for font_path in font_paths:
            try:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                    font_name = 'ChineseFont'
                    font_registered = True
                    st.success(f"成功加载字体: {font_path}")
                    break
            except Exception as e:
                st.warning(f"字体加载失败 {font_path}: {e}")
                continue
        
        # 方法2: 尝试使用系统字体 (适用于Streamlit Cloud的Linux环境)
        if not font_registered:
            system_fonts = [
                '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
                '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            ]
            
            for font_path in system_fonts:
                try:
                    if os.path.exists(font_path):
                        pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                        font_name = 'ChineseFont'
                        font_registered = True
                        st.success(f"成功加载系统字体: {font_path}")
                        break
                except Exception as e:
                    st.warning(f"系统字体加载失败 {font_path}: {e}")
                    continue
        
        # 方法3: 使用Base64编码的字体作为备用方案
        if not font_registered:
            # 这里可以使用Base64编码的字体数据
            # 例如，您可以预先将字体文件转换为Base64字符串
            font_data_base64 = None  # 这里应该是您的Base64编码字体数据
            
            if font_data_base64:
                try:
                    # 创建临时字体文件
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.ttf') as tmp:
                        tmp.write(base64.b64decode(font_data_base64))
                        font_path = tmp.name
                    
                    # 注册字体
                    pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                    font_name = 'ChineseFont'
                    font_registered = True
                    st.success("成功加载Base64字体")
                except Exception as e:
                    st.warning(f"Base64字体加载失败: {e}")
    
    except Exception as e:
        st.error(f"字体设置过程中发生错误: {e}")
        font_name = 'Helvetica'  # 回退到默认字体
    
    # 创建中文字体样式
    if font_name != 'Helvetica':
        # 添加中文样式
        styles.add(ParagraphStyle(
            name='ChineseTitle',
            parent=styles['Title'],
            fontName=font_name,
            fontSize=16,
            spaceAfter=30,
        ))
        
        styles.add(ParagraphStyle(
            name='ChineseHeading',
            parent=styles['Heading1'],
            fontName=font_name,
            fontSize=14,
            spaceAfter=12,
        ))
        
        styles.add(ParagraphStyle(
            name='ChineseBody',
            parent=styles['BodyText'],
            fontName=font_name,
            fontSize=10,
            spaceAfter=6,
        ))
        
        # 使用中文样式
        title_style = styles['ChineseTitle']
        heading_style = styles['ChineseHeading']
        body_style = styles['ChineseBody']
    else:
        # 使用默认样式
        title_style = styles['Title']
        heading_style = styles['Heading1']
        body_style = styles['BodyText']
        st.warning("使用默认字体Helvetica，中文可能无法正确显示")
    
    # 创建样式
    styles = getSampleStyleSheet()
    
    # 自定义标题样式
    title_style = ParagraphStyle(
        'CustomTitle',
        fontName=font_name,
        fontSize=16,
        alignment=TA_CENTER,
        spaceAfter=20
    )
    
    # 自定义子标题样式
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        fontName=font_name,
        fontSize=14,
        textColor=colors.black,
        spaceAfter=10
    )
    
    # 自定义正文样式 - 添加中文排版支持
    body_style = ParagraphStyle(
        'CustomBody',
        fontName=font_name,
        fontSize=12,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=10,
        wordWrap='CJK'  # 添加中文换行支持
    )
    
    # 自定义信息样式
    info_style = ParagraphStyle(
        'CustomInfo',
        fontName=font_name,
        fontSize=12,
        textColor=colors.grey,
        spaceAfter=5
    )
    
    # 构建内容
    story = []
    
    # 标题
    story.append(Paragraph('电机故障诊断报告', title_style))
    story.append(Spacer(1, 20))
    
    # 诊断信息
    story.append(Paragraph('诊断结果', subtitle_style))
    story.append(Paragraph(f'故障类型: {diagnosis_class}', info_style))
    story.append(Paragraph(f'置信度: {confidence_level*100:.2f}%', info_style))
    story.append(Spacer(1, 15))
    
    # 分析结果
    story.append(Paragraph('信号分析', subtitle_style))
    
    # 创建分析结果表格
    analysis_data = [['参数', '值']]
    for key, value in analysis_results.items():
        analysis_data.append([key, str(value)])
    
    analysis_table = Table(analysis_data, colWidths=[2*inch, 3*inch])
    analysis_table.setStyle(TableStyle([
        ('FONT', (0, 0), (-1, -1), font_name, 10),
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONT', (0, 0), (-1, 0), font_name, 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(analysis_table)
    story.append(Spacer(1, 20))
    
    # 添加AI分析标题
    ai_title_style = ParagraphStyle(
        'AITitle',
        fontName=font_name,
        fontSize=14,
        textColor=colors.black,
        spaceAfter=10,
        spaceBefore=20
    )
    story.append(Paragraph('AI分析与建议', ai_title_style))
    
    # 处理LLM报告内容
    # 将报告内容按行分割，处理列表项
    report_lines = llm_report.split('\n')
    
    # 用于跟踪列表状态
    in_list = False
    list_items = []
    list_counter = 1  # 列表项计数器
    
    def flush_list(reset_counter=False):
        """将当前列表项添加到story中"""
        nonlocal list_items, list_counter
        if not list_items:
            return
            
        # 创建表格来显示列表项，确保序号和内容在同一行
        table_data = []
        for i, item in enumerate(list_items, list_counter):
            # 处理加粗文本
            item = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', item)
            
            # 将序号和内容放在同一行
            table_data.append([
                Paragraph(f"({i})", body_style),
                Paragraph(item, body_style)
            ])
        
        # 创建表格 - 使用紧凑布局
        list_table = Table(table_data, colWidths=[0.3*inch, 6.2*inch])
        list_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
            ('LEFTPADDING', (0, 0), (0, -1), 1),
            ('RIGHTPADDING', (0, 0), (0, -1), 1),
            ('LEFTPADDING', (1, 0), (1, -1), 1),
            ('RIGHTPADDING', (1, 0), (1, -1), 1),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ('TOPPADDING', (0, 0), (-1, -1), 2),
        ]))
        
        story.append(list_table)
        story.append(Spacer(1, 4))
        
        list_items = []
        
        # 如果需要重置计数器，则重置为1
        if reset_counter:
            list_counter = 1
        else:
            # 否则更新计数器
            list_counter += len(table_data)
    
    # 文本清理函数
    def clean_text(text):
        """清理文本，确保不以标点符号开头"""
        if not text:
            return text
        
        # 移除开头的标点符号
        text = re.sub(r'^[，。！？；：,.!?;:\s]+', '', text)
        
        # 移除多余的空格
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    # 中文文本处理函数
    def process_chinese_text(text):
        """处理中文文本，优化排版"""
        if not text:
            return text
        
        # 处理常见的标点符号连接问题
        # 确保标点符号跟随在前面的文本后面
        text = re.sub(r'([^，。！？；：])\s*([，。！？；：])', r'\1\2', text)
        
        # 处理英文和数字与中文的间距
        text = re.sub(r'([a-zA-Z0-9])\s*([\u4e00-\u9fff])', r'\1 \2', text)
        text = re.sub(r'([\u4e00-\u9fff])\s*([a-zA-Z0-9])', r'\1 \2', text)
        
        return text
    
    # 段落缓冲区，用于合并多行文本形成完整段落
    paragraph_buffer = []
    
    def flush_paragraph():
        """将缓冲区中的文本合并为一个段落并添加到story中"""
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        
        # 合并缓冲区中的文本
        paragraph_text = ''.join(paragraph_buffer)  # 使用空字符串连接，避免添加额外空格
        paragraph_text = clean_text(paragraph_text)
        paragraph_text = process_chinese_text(paragraph_text)  # 处理中文排版
        
        if paragraph_text:
            # 处理加粗文本
            paragraph_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', paragraph_text)
            story.append(Paragraph(paragraph_text, body_style))
            story.append(Spacer(1, 6))
        
        paragraph_buffer = []
    
    for line in report_lines:
        line = line.strip()
        if not line:
            # 空行表示段落结束
            flush_paragraph()
            if in_list:
                flush_list()
                in_list = False
            continue
            
        # 检查是否是列表项
        is_list_item = line.startswith('- ') or line.startswith('* ') or re.match(r'^\d+\.\s+', line)
        
        if is_list_item:
            # 刷新当前段落
            flush_paragraph()
            
            # 开始或继续列表
            in_list = True
            
            # 提取列表项内容
            if line.startswith('- ') or line.startswith('* '):
                list_item = line[2:].strip()  # 移除 "- " 或 "* "
            else:
                # 对于有序列表，移除数字和点
                list_item = re.sub(r'^\d+\.\s+', '', line)
                
            # 清理列表项文本
            list_item = clean_text(list_item)
            list_item = process_chinese_text(list_item)  # 处理中文排版
            if list_item:
                list_items.append(list_item)
        else:
            # 如果不是列表项，刷新当前列表
            if in_list:
                flush_list()
                in_list = False
            
            # 检查是否是标题
            is_title = False
            if line.startswith('#### '):
                # 四级标题
                is_title = True
                flush_paragraph()  # 刷新当前段落
                title_text = line[5:].strip()
                title_text = clean_text(title_text)
                title_text = process_chinese_text(title_text)  # 处理中文排版
                title_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', title_text)
                title_style = ParagraphStyle(
                    'ReportTitle4',
                    fontName=font_name,
                    fontSize=11,
                    textColor=colors.black,
                    spaceAfter=6,
                    spaceBefore=12
                )
                story.append(Paragraph(title_text, title_style))
            elif line.startswith('### '):
                # 三级标题
                is_title = True
                flush_paragraph()  # 刷新当前段落
                title_text = line[4:].strip()
                title_text = clean_text(title_text)
                title_text = process_chinese_text(title_text)  # 处理中文排版
                title_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', title_text)
                title_style = ParagraphStyle(
                    'ReportTitle3',
                    fontName=font_name,
                    fontSize=12,
                    textColor=colors.black,
                    spaceAfter=6,
                    spaceBefore=12
                )
                story.append(Paragraph(title_text, title_style))
            elif line.startswith('## '):
                # 二级标题
                is_title = True
                flush_paragraph()  # 刷新当前段落
                title_text = line[3:].strip()
                title_text = clean_text(title_text)
                title_text = process_chinese_text(title_text)  # 处理中文排版
                title_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', title_text)
                title_style = ParagraphStyle(
                    'ReportTitle2',
                    fontName=font_name,
                    fontSize=14,
                    textColor=colors.black,
                    spaceAfter=6,
                    spaceBefore=12
                )
                story.append(Paragraph(title_text, title_style))
            elif line.startswith('# '):
                # 一级标题
                is_title = True
                flush_paragraph()  # 刷新当前段落
                title_text = line[2:].strip()
                title_text = clean_text(title_text)
                title_text = process_chinese_text(title_text)  # 处理中文排版
                title_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', title_text)
                title_style = ParagraphStyle(
                    'ReportTitle1',
                    fontName=font_name,
                    fontSize=16,
                    textColor=colors.black,
                    spaceAfter=6,
                    spaceBefore=12
                )
                story.append(Paragraph(title_text, title_style))
            elif line.startswith('---') or line.startswith('***'):
                # 分割线
                flush_paragraph()  # 刷新当前段落
                story.append(Spacer(1, 12))
                hr = Table([['']], colWidths=[6*inch], rowHeights=[1])
                hr.setStyle(TableStyle([
                    ('LINEABOVE', (0, 0), (-1, -1), 1, colors.grey),
                ]))
                story.append(hr)
                story.append(Spacer(1, 12))
            else:
                # 普通段落 - 添加到缓冲区
                line = process_chinese_text(line)  # 处理中文排版
                paragraph_buffer.append(line)
            
            # 如果是标题，重置列表计数器
            if is_title:
                list_counter = 1
    
    # 处理最后可能还在缓冲区或列表中的内容
    flush_paragraph()
    if in_list:
        flush_list()
    
    # 添加页脚
    story.append(Spacer(1, 20))
    footer_style = ParagraphStyle(
        'CustomFooter',
        fontName=font_name,
        fontSize=10,
        alignment=TA_CENTER,
        textColor=colors.grey
    )
    story.append(Paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", footer_style))
    
    # 构建PDF
    doc.build(story)
    
    # 获取PDF数据
    pdf_data = buffer.getvalue()
    buffer.close()
    
    return pdf_data

def generate_diagnostic_report(diagnosis_class, confidence_level, analysis_results):
    # 检查是否已经存在诊断结果，如果不存在则设置
    if 'diagnosis_results' not in st.session_state or st.session_state['diagnosis_results'] is None:
        st.session_state['diagnosis_results'] = {
            'diagnosis_class': diagnosis_class,
            'confidence_level': confidence_level,
            'analysis_results': analysis_results,
            'report': None,
            'pdf_data': None,  # 新增：保存PDF数据
            'docx_data': None  # 新增：保存DOCX数据
        }
    else:
        # 如果已经存在，则更新诊断结果（确保是最新的）
        st.session_state['diagnosis_results'].update({
            'diagnosis_class': diagnosis_class,
            'confidence_level': confidence_level,
            'analysis_results': analysis_results
        })

    # 准备故障手册信息 - 使用向量知识库检索相关内容
    manual_info = ""
    if ('vector_index' in st.session_state and 
        'vector_chunks' in st.session_state and 
        st.session_state['vector_index'] is not None and
        st.session_state['vector_chunks'] is not None):
        
        # 从向量知识库中检索与诊断结果相关的内容
        embedding_model = load_embedding_model()
        if embedding_model:
            # 使用诊断结果作为查询
            query = f"{diagnosis_class} {analysis_results}"
            relevant_chunks = search_vector_knowledge_base(
                query, 
                st.session_state['vector_index'],
                st.session_state['vector_chunks'],
                embedding_model,
                top_k=3
            )
            
            if relevant_chunks:
                manual_info = "以下是根据故障手册检索到的相关信息：\n"
                for i, chunk in enumerate(relevant_chunks):
                    manual_info += f"\n[相关段落 {i+1}, 相似度: {chunk['score']:.3f}]\n{chunk['text']}\n"
    
    llm_prompt = f"""
    已知振动信号经过CNN模型诊断，结果如下：
    - 诊断结果: {diagnosis_class}
    - 置信度: {confidence_level*100:.2f}%
    - 信号统计: {analysis_results}

    {manual_info}

    请你作为机械设备故障诊断专家，根据上述诊断结果和故障手册参考，生成一份专业、详细的诊断报告，报告内容应包括：
    1. 故障解释：详细说明为什么可能是该故障，从机械原理角度解释
    2. 可能的成因：分析导致此类故障的多种可能原因
    3. 建议的处理措施：提供具体、可行的维修和处理建议，按优先级排序
    4. 预防措施：建议如何预防此类故障再次发生

    请确保报告专业、准确且易于理解，使用中文回答。
    """
    
    st.subheader("📊 AI大模型解释与建议")
    
    # 将字符串转换为消息格式
    messages = [{"role": "user", "content": llm_prompt}]

    report = call_ai_api(messages)
    
    if report:
        st.success("报告生成完成！")
        st.write(report)

        # 更新session_state中的报告内容
        st.session_state['diagnosis_results']['report'] = report
        
        try:
            # 创建PDF报告并保存数据
            pdf_data = create_pdf_report(
                diagnosis_class, 
                confidence_level, 
                analysis_results,
                report
            )
            st.session_state['diagnosis_results']['pdf_data'] = pdf_data
            
            # 创建DOCX报告并保存数据
            docx_data = create_docx_report(
                diagnosis_class, 
                confidence_level, 
                analysis_results,
                report
            )
            st.session_state['diagnosis_results']['docx_data'] = docx_data
            
            # 显示下载按钮
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="下载诊断报告 (PDF)",
                    data=pdf_data,
                    file_name=f"设备诊断报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf"
                )
            with col2:
                st.download_button(
                    label="下载诊断报告 (DOCX)",
                    data=docx_data,
                    file_name=f"设备诊断报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
        except Exception as e:
            st.error(f"生成报告文件时出错: {e}")
            st.info("如果您需要PDF或DOCX版本，请确保系统安装了必要的库和字体")
    else:
        st.error("生成报告失败，请检查API配置或稍后重试")

# 初始化session state变量
def initialize_session_state():
    """初始化所有session state变量"""
    if 'diagnosis_completed' not in st.session_state:
        st.session_state.diagnosis_completed = False
    if 'predicted_class' not in st.session_state:
        st.session_state.predicted_class = None
    if 'confidence_value' not in st.session_state:
        st.session_state.confidence_value = None
    if 'analysis_results' not in st.session_state:
        st.session_state.analysis_results = {}
    if 'show_report' not in st.session_state:
        st.session_state.show_report = False
    if 'diagnosis_results' not in st.session_state:
        st.session_state.diagnosis_results = None
    if "messages" not in st.session_state:
        st.session_state["messages"] = [
            {"role": "system", "content": "你是一名电机故障诊断专家，帮助用户分析信号并给出诊断建议。"}
        ]
    if 'manual_content' not in st.session_state:
        st.session_state.manual_content = None
    if 'vector_index' not in st.session_state:
        st.session_state.vector_index = None
    if 'vector_chunks' not in st.session_state:
        st.session_state.vector_chunks = None
    if 'processing_mode' not in st.session_state:
        st.session_state.processing_mode = "自动"  # 默认值
    if 'avg_probabilities' not in st.session_state:
        st.session_state.avg_probabilities = None
    if 'preferred_api' not in st.session_state:
        st.session_state.preferred_api = "deepseek"  # 默认优先使用DeepSeek

# 调用初始化函数
initialize_session_state()

# 调用加载模型函数
model = load_model()

# 添加侧边栏信息
with st.sidebar:
    st.header("关于")
    st.markdown("""
    本系统基于深度学习CNN模型，用于电机轴承故障诊断。
    
    ### 支持诊断的故障类型:
    - 正常状态
    - 7mil内圈故障
    - 7mil滚动体故障
    - 7mil外圈故障
    - 14mil内圈故障
    - 14mil滚动体故障
    - 14mil外圈故障
    - 21mil内圈故障
    - 21mil滚动体故障
    - 21mil外圈故障
    
    ### 使用方法:
    1. 上传CSV格式的振动数据文件
    2. 选择要分析的数据列(如果有多个列)
    3. 点击"开始故障诊断"按钮
    4. 查看诊断结果和建议
    """)
    
    st.header("模型信息")
    if model is not None:
        st.success("模型已加载")
        st.write(f"设备: {'GPU' if device.type == 'cuda' else 'CPU'}")
    else:
        st.error("模型未加载")

# 侧边栏设置
with st.sidebar:
    st.header("API设置")
    
    # 添加API选择选项
    api_option = st.radio(
        "选择优先使用的API",
        ["DeepSeek", "KIMI"],
        index=0 if st.session_state.preferred_api == "deepseek" else 1,
        help="选择优先使用哪个API来回答问题"
    )
    
    # 更新首选的API
    st.session_state.preferred_api = "deepseek" if api_option == "DeepSeek" else "kimi"
    
    st.header("性能设置")
    
    # 添加性能配置选项
    chunk_size = st.slider("文本块大小", min_value=200, max_value=1000, value=500, step=100,
                         help="较小的值提高精度但增加处理时间，较大的值减少处理时间但可能降低精度")
    
    overlap_size = st.slider("重叠大小", min_value=0, max_value=200, value=50, step=10,
                           help="文本块之间的重叠字符数，有助于保持上下文完整性")
    
    # 添加参数验证
    if overlap_size >= chunk_size:
        st.error("错误: 重叠大小必须小于文本块大小")
        # 自动调整重叠大小
        overlap_size = min(overlap_size, chunk_size - 10)
        st.info(f"已自动调整重叠大小为: {overlap_size}")
    
    batch_size = st.slider("批量处理大小", min_value=8, max_value=64, value=16, step=8,
                         help="较大的批量提高处理速度但需要更多内存")
    
    st.header("API测试")
    
    # 添加DeepSeek API测试按钮
    if st.button("测试DeepSeek API连接"):
        test_prompt = "你好，请回复'连接成功'以确认API连接正常。"
        test_messages = [{"role": "user", "content": test_prompt}]

        # 临时清除诊断结果，避免影响测试
        original_diagnosis = st.session_state.get('diagnosis_results')
        st.session_state['diagnosis_results'] = None
        
        with st.spinner("测试DeepSeek API连接中..."):
            res = call_deepseek_api(test_messages, max_tokens=50)
        
        # 恢复原始诊断结果
        st.session_state['diagnosis_results'] = original_diagnosis

        if res and "连接成功" in res:
            st.success(f"DeepSeek API连接测试成功！响应: {res}")
        else:
            st.error(f"DeepSeek API连接测试失败: {res}")
    
    # 添加KIMI API测试按钮
    if st.button("测试KIMI API连接"):
        test_prompt = "你好，请回复'连接成功'以确认API连接正常。"
        test_messages = [{"role": "user", "content": test_prompt}]

        # 临时清除诊断结果，避免影响测试
        original_diagnosis = st.session_state.get('diagnosis_results')
        st.session_state['diagnosis_results'] = None
        
        with st.spinner("测试KIMI API连接中..."):
            res = call_kimi_api(test_messages, max_tokens=50)
        
        # 恢复原始诊断结果
        st.session_state['diagnosis_results'] = original_diagnosis

        if res and "连接成功" in res:
            st.success(f"KIMI API连接测试成功！响应: {res}")
        else:
            st.error(f"KIMI API连接测试失败: {res}")
    
    st.header("故障手册设置")
    manual_file = st.file_uploader("上传故障手册 (PDF/TXT/DOCX)", type=["pdf", "txt", "docx"])
    
    if manual_file is not None:
        # 提取文本内容
        manual_text = extract_text_from_file(manual_file)
        if manual_text:
            # 预处理文本
            manual_text = preprocess_text(manual_text)
            
            st.session_state.manual_content = manual_text
            
            # 显示文本统计信息
            text_length = len(manual_text)
            
            # 修复分块计算逻辑
            if chunk_size <= overlap_size:
                st.error("错误: 文本块大小必须大于重叠大小")
                estimated_chunks = 1
            else:
                # 计算预计分块数 - 修复后的公式
                if text_length <= chunk_size:
                    estimated_chunks = 1
                else:
                    estimated_chunks = max(1, (text_length - overlap_size) // (chunk_size - overlap_size) + 1)
            
            st.info(f"文本长度: {text_length} 字符, 预计分成: {estimated_chunks} 个文本块")
            
            # 根据处理模式决定是否创建向量索引
            if st.session_state.processing_mode == "仅文本搜索" or (st.session_state.processing_mode == "自动" and text_length > 200000):
                # 对于长文本，使用简单文本搜索
                chunks = []
                start = 0
                while start < len(manual_text):
                    end = min(start + chunk_size, len(manual_text))
                    chunk = manual_text[start:end]
                    chunks.append(chunk)
                    start = end - overlap_size  # 考虑重叠
                
                st.session_state.vector_index = None
                st.session_state.vector_chunks = chunks
                st.info(f"已创建 {len(chunks)} 个文本片段用于简单搜索")
            else:
                # 创建向量知识库（使用优化版本）
                embedding_model = load_embedding_model()
                if embedding_model:
                    with st.spinner("正在构建向量知识库..."):
                        index, chunks = create_vector_knowledge_base_optimized(
                            manual_text, 
                            embedding_model,
                            chunk_size=chunk_size,
                            overlap=overlap_size,
                            batch_size=batch_size
                        )
                        st.session_state.vector_index = index
                        st.session_state.vector_chunks = chunks
                    
                    if index is not None:
                        st.success(f"故障手册上传成功! 已构建 {len(chunks)} 个知识片段")
                    else:
                        st.warning("使用简单文本搜索模式，某些高级功能可能受限")
                else:
                    st.error("嵌入模型加载失败，使用简单文本搜索模式")
                    # 即使没有嵌入模型，也创建简单的文本库
                    chunks = []
                    start = 0
                    while start < len(manual_text):
                        end = min(start + chunk_size, len(manual_text))
                        chunk = manual_text[start:end]
                        chunks.append(chunk)
                        start = end - overlap_size  # 考虑重叠
                    
                    st.session_state.vector_index = None
                    st.session_state.vector_chunks = chunks
                    st.info(f"已创建 {len(chunks)} 个文本片段用于简单搜索")
        else:
            st.error("提取文本失败，请重新上传文件。")
    elif 'manual_content' not in st.session_state:
        st.session_state.manual_content = None
        st.session_state.vector_index = None
        st.session_state.vector_chunks = None
    
    # 显示当前状态
    if st.session_state.vector_index is not None:
        st.info(f"已加载故障手册，包含 {len(st.session_state.vector_chunks)} 个知识片段")
        if st.button("清除故障手册", key="clear_manual_btn"):
            # 彻底清除所有与手册相关的状态
            st.session_state.manual_content = None
            st.session_state.vector_index = None
            st.session_state.vector_chunks = None
            
            # 清除嵌入模型缓存（重要！）
            if 'load_embedding_model' in st.session_state:
                del st.session_state['load_embedding_model']
            
            # 清除诊断结果中的手册引用
            if 'diagnosis_results' in st.session_state:
                # 保留诊断结果但移除手册内容
                diagnosis = st.session_state['diagnosis_results']
                if diagnosis and 'manual_references' in diagnosis:
                    del diagnosis['manual_references']
            
            # 清除对话历史中可能包含的手册内容
            st.session_state["messages"] = [
                {"role": "system", "content": "你是一名电机故障诊断专家，帮助用户分析信号并给出诊断建议。"}
            ]
            
            st.success("故障手册已彻底清除!")
            st.rerun()
    else:
        st.warning("未上传故障手册")

# 在侧边栏添加清除缓存按钮
st.sidebar.header("高级设置")
if st.sidebar.button("强制清除所有缓存", key="clear_all_cache_btn"):
    # 清除所有缓存和会话状态
    keys_to_keep = []  # 保留必要的键
    keys_to_delete = [key for key in st.session_state.keys() if key not in keys_to_keep]
    
    for key in keys_to_delete:
        del st.session_state[key]
    
    # 清除资源缓存
    if 'load_embedding_model' in st.session_state:
        del st.session_state['load_embedding_model']
    
    st.sidebar.success("所有缓存已清除!")
    st.rerun()

# 主界面
st.header("振动数据分析")

# 文件上传区域
uploaded_file = st.file_uploader("选择振动数据文件 (CSV格式)", type=["csv"])

if uploaded_file is not None and model is not None:
    try:
        dataframe = pd.read_csv(uploaded_file)
        st.success("文件上传成功!")
        
        # 显示数据基本信息
        st.subheader("数据概览")
        col1, col2, col3 = st.columns(3)
        col1.metric("数据行数", dataframe.shape[0])
        col2.metric("数据列数", dataframe.shape[1])
        col3.metric("采样点数", dataframe.shape[0])
        
        # 选择要分析的数据列
        if dataframe.shape[1] > 1:
            selected_column = st.selectbox("选择要分析的数据列", dataframe.columns)
            vibration_data = dataframe[selected_column].values
        else:
            vibration_data = dataframe.iloc[:, 0].values
        
        # 显示数据波形
        st.subheader("振动信号波形")
        fig, ax = plt.subplots(figsize=(10, 4))
        ax.plot(vibration_data[:1024])
        ax.set_xlabel("采样点")
        ax.set_ylabel("振幅")
        ax.set_title("振动信号波形 (前1024个点)")
        ax.grid(True)
        st.pyplot(fig)
        
        # 诊断按钮
        if st.button("开始故障诊断", type="primary", key="diagnose_btn"):
            if len(vibration_data) < 1024:
                st.error(f"数据长度不足1024点，当前只有{len(vibration_data)}点")
            else:
                with st.spinner("正在分析数据..."):
                    # 滑动窗口分割
                    time_steps = 1024
                    overlap_ratio = 0.5
                    split_new_data = split_data_with_overlap(
                        vibration_data, time_steps, label=-1, overlap_ratio=overlap_ratio
                    )
                    
                    # 提取特征部分并转换为张量
                    input_new = split_new_data.iloc[:, 0:time_steps]
                    input_new_tensor = torch.tensor(input_new.values).float().to(device)

                    # 模型预测
                    predictions = []
                    all_probabilities_list = []

                    with torch.no_grad():
                        for i in range(input_new_tensor.size(0)):
                            output = model(input_new_tensor[i].unsqueeze(0))
                            _, predicted = torch.max(output.data, 1)
                            predictions.append(predicted.item())
                            probabilities = torch.nn.functional.softmax(output, dim=1)
                            all_probabilities_list.append(probabilities.cpu().numpy()[0])

                    # 计算整体预测结果
                    if predictions:
                        prediction_counter = Counter(predictions)
                        most_common = prediction_counter.most_common(1)
                        predicted_label = most_common[0][0] if most_common else 0
                        predicted_class = fault_type_mapping.get(predicted_label, f"未知故障({predicted_label})")
                        
                        avg_probabilities = np.mean(all_probabilities_list, axis=0)
                        confidence_value = avg_probabilities[predicted_label]
                        
                        # 存储分析结果
                        analysis_results = {
                            "最小值": f"{np.min(vibration_data):.4f}",
                            "最大值": f"{np.max(vibration_data):.4f}",
                            "平均值": f"{np.mean(vibration_data):.4f}",
                            "标准差": f"{np.std(vibration_data):.4f}",
                            "数据点数": len(vibration_data),
                            "分析样本数": len(predictions)
                        }
                        
                        # 存储结果到session state
                        st.session_state.diagnosis_completed = True
                        st.session_state.predicted_class = predicted_class
                        st.session_state.confidence_value = confidence_value
                        st.session_state.analysis_results = analysis_results
                        st.session_state.avg_probabilities = avg_probabilities
                        st.session_state.show_report = False  # 重置报告显示状态

                        # 立即设置diagnosis_results，以便在对话中使用
                        st.session_state['diagnosis_results'] = {
                            'diagnosis_class': predicted_class,
                            'confidence_level': confidence_value,
                            'analysis_results': analysis_results,
                            'report': None  # 此时还没有报告，生成报告后再更新
                        }
                        
                        # 添加一点延迟
                        time.sleep(1)
    
    except Exception as e:
        st.error(f"文件处理错误: {e}")

# 显示诊断结果（如果已完成诊断）
if st.session_state.diagnosis_completed:
    st.markdown("---")
    st.subheader("🔍 诊断结果")
    
    # 添加提示信息
    st.info("诊断已完成！您现在可以与AI助手对话，获取基于诊断结果的详细解释和建议。")
    
    # 使用列布局显示主要结果
    col1, col2 = st.columns(2)
    
    with col1:
        if st.session_state.predicted_class.startswith("正常状态"):
            st.success(f"**诊断结果**: {st.session_state.predicted_class}")
        else:
            st.error(f"**诊断结果**: {st.session_state.predicted_class}")
        st.info(f"**置信度**: {st.session_state.confidence_value*100:.2f}%")
    
    with col2:
        # 显示概率最高的前3个故障类型
        top3_indices = np.argsort(st.session_state.avg_probabilities)[-3:][::-1]
        st.write("**概率最高的故障类型**:")
        for i, idx in enumerate(top3_indices):
            prob = st.session_state.avg_probabilities[idx] * 100
            fault_name = fault_type_mapping.get(idx, f"未知故障{idx}")
            st.write(f"{i+1}. {fault_name}: {prob:.2f}%")
    
    # 显示故障概率分布
    st.subheader("故障概率分布")
    prob_df = pd.DataFrame({
        "故障类型": [fault_type_mapping.get(i, f"未知{i}") for i in range(10)],
        "概率(%)": [p * 100 for p in st.session_state.avg_probabilities]
    })
    st.bar_chart(prob_df.set_index("故障类型"))
    
    # 显示详细分析
    with st.expander("查看详细分析"):
        st.write("**原始振动信号统计**:")
        for key, value in st.session_state.analysis_results.items():
            st.write(f"- {key}: {value}")
        
        st.write("**分析说明**:")
        if st.session_state.predicted_class.startswith("正常状态"):
            st.write("振动信号特征显示电机处于正常运行状态，无明显故障特征。")
        else:
            st.write(f"检测到故障特征，最可能的原因是: {st.session_state.predicted_class.split('(')[0]}")
            
            if "内圈" in st.session_state.predicted_class:
                st.write("**建议**: 检查电机轴承内圈是否有磨损、裂纹或点蚀。")
            elif "外圈" in st.session_state.predicted_class:
                st.write("**建议**: 检查电机轴承外圈是否有磨损、裂纹或点蚀。")
            elif "滚动体" in st.session_state.predicted_class:
                st.write("**建议**: 检查电机轴承滚动体是否有磨损、裂纹或缺失。")
            
            st.write("建议尽快安排专业人员进行检查和维护。")
    
    # 生成诊断报告按钮
    if st.button("生成诊断报告", key="generate_report_btn"):
        st.session_state.show_report = True
    
    # # 显示诊断报告（如果用户点击了生成报告按钮）
    # if st.session_state.show_report:
    #     generate_diagnostic_report(
    #         st.session_state.predicted_class, 
    #         st.session_state.confidence_value,
    #         st.session_state.analysis_results
    #     )

    # 显示诊断报告（如果用户点击了生成报告按钮）
if st.session_state.show_report:
    # 检查是否已经有生成的报告数据
    if (st.session_state['diagnosis_results'] and 
        st.session_state['diagnosis_results'].get('pdf_data') and 
        st.session_state['diagnosis_results'].get('docx_data')):
        
        # 直接显示下载按钮，不重新生成报告
        st.subheader("📊 AI大模型解释与建议")
        st.write(st.session_state['diagnosis_results']['report'])
        
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                label="下载诊断报告 (PDF)",
                data=st.session_state['diagnosis_results']['pdf_data'],
                file_name=f"设备诊断报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                mime="application/pdf"
            )
        with col2:
            st.download_button(
                label="下载诊断报告 (DOCX)",
                data=st.session_state['diagnosis_results']['docx_data'],
                file_name=f"设备诊断报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        # 如果没有生成的报告数据，调用函数生成
        generate_diagnostic_report(
            st.session_state.predicted_class, 
            st.session_state.confidence_value,
            st.session_state.analysis_results
        )

    # 显示诊断状态
    if 'diagnosis_results' in st.session_state and st.session_state.diagnosis_results:
        st.success("✅ 已有诊断结果，AI助手将基于此结果回答您的问题")
    else:
        st.info("ℹ️ 暂无诊断结果，AI助手将提供一般性建议。请先生成诊断报告。")

    # 清除按钮
    if st.button("清除诊断结果和对话历史"):
        if 'diagnosis_results' in st.session_state:
            del st.session_state['diagnosis_results']
        st.session_state["messages"] = [
            {"role": "system", "content": "你是一名电机故障诊断专家，帮助用户分析信号并给出诊断建议。"}
        ]
        st.success("已清除诊断结果和对话历史")
        st.rerun()

    # 聊天框提示
    if st.session_state.vector_index is not None:
        st.info("📚 当前已加载故障手册向量知识库，AI助手将智能检索相关内容")
    else:
        st.warning("📚 未加载故障手册，AI助手将基于通用知识回答。您可以在侧边栏上传故障手册")

    # -----------------------------
    # 聊天框 - 放在诊断结果下方
    # -----------------------------
    st.markdown("---")
    st.subheader("💬 电机诊断助手对话")

    # 显示历史对话
    for msg in st.session_state["messages"]:
        if msg["role"] == "user":
            st.chat_message("user").write(msg["content"])
        elif msg["role"] == "assistant":
            st.chat_message("assistant").write(msg["content"])

    # 输入框
    if prompt := st.chat_input("输入问题，例如：请分析当前电机信号是否异常"):
        # 保存用户输入
        st.session_state["messages"].append({"role": "user", "content": prompt})
        st.chat_message("user").write(prompt)

        # -----------------------------
        # 调用AI API接口
        # -----------------------------
        try:
            reply = call_ai_api(
                messages=st.session_state["messages"],
                max_tokens=2000,
                temperature=0.7
            )
            
            if reply is None:
                reply = "❌ 模型调用失败，请查看错误信息"
                
        except Exception as e:
            reply = f"❌ 模型调用失败: {str(e)}"

        # 保存助手回复 
        st.session_state["messages"].append({"role": "assistant", "content": reply})
        st.chat_message("assistant").write(reply)
# 只在没有上传文件且没有诊断结果时显示提示
if uploaded_file is None and not st.session_state.diagnosis_completed:
    st.info("请上传CSV格式的振动数据文件开始分析")

# 添加页脚
st.markdown("---")

st.markdown("电机故障诊断系统 🔧 2025 | 基于PyTorch和Streamlit开发")
